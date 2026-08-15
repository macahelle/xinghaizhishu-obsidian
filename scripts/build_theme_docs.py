from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FONT = "PingFang SC"
INK = "172033"
BLUE = "38318B"
CYAN = "18B9CA"
MUTED = "68748A"
FILL = "E8EEF5"
WARN = "FFF1ED"


def set_font(run, size=11, bold=False, color=INK):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), FONT)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if shd.getparent() is None:
        tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")
    if tc_w.getparent() is None:
        tc_pr.append(tc_w)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)


def configure(doc, short_title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for name, size, before, after in (("Heading 1", 16, 18, 10), ("Heading 2", 13, 14, 7), ("Heading 3", 12, 10, 5)):
        style = styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    header = section.header.paragraphs[0]
    run = header.add_run(f"星海知枢  |  {short_title}")
    set_font(run, 8.5, True, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("主题 3.2.0  ·  工作台 1.2.6")
    set_font(run, 8.5, False, MUTED)


def title(doc, text, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(55)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("星海知枢"), 14, True, CYAN)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(text), 27, True, BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    set_font(p.add_run(subtitle), 13, False, MUTED)
    add_table(doc, [
        ("产品类型", "Obsidian 主题与配套工作台插件"),
        ("主题版本", "3.2.0"),
        ("插件版本", "1.2.6"),
        ("最低版本", "Obsidian 1.9.0"),
        ("适用系统", "macOS 与 Windows"),
        ("更新日期", "2026-08-15"),
    ])
    note(doc, "范围说明", "发布包包含主题和 xinghai-workbench 插件，不包含用户 data.json、测试库、QA 证据或归档数据。")
    doc.add_page_break()


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "说明"
    widths = (2700, 6660)
    set_table_geometry(table, widths)
    for row_index, row in enumerate(table.rows):
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            shade(cell, FILL)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                set_font(run, 10, True, BLUE)
    for label, detail in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
        for index, cell in enumerate(cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                set_font(run, 10, index == 0, INK)
    doc.add_paragraph()


def note(doc, label, text, warning=False):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_width(cell, 9360)
    set_cell_margins(cell)
    set_table_geometry(table, (9360,))
    shade(cell, WARN if warning else "EEF2FF")
    p = cell.paragraphs[0]
    set_font(p.add_run(f"{label}："), 10, True, "A33424" if warning else BLUE)
    set_font(p.add_run(text), 10)
    doc.add_paragraph()


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def body(doc, text):
    p = doc.add_paragraph()
    set_font(p.add_run(text))


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            set_font(run)
        if not p.runs:
            set_font(p.add_run(item))
        else:
            p.runs[0].text = item


def steps(doc, items):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(f"{index}.  {item}"))


def build_install():
    doc = Document()
    configure(doc, "macOS / Windows 安装指导")
    title(doc, "macOS / Windows 安装指导说明", "从下载、复制到启用与回滚")
    heading(doc, "1. 安装前确认")
    bullets(doc, [
        "本包包含 Obsidian 主题 3.2.0 和工作台插件 1.2.6，不是独立应用。",
        "确认 Obsidian 版本不低于 1.9.0。",
        "建议备份知识库中的 `.obsidian/appearance.json`。",
        "升级用户另行备份 `.obsidian/plugins/xinghai-workbench/data.json`。",
        "安装和覆盖文件前完全退出 Obsidian。",
    ])
    heading(doc, "2. 安装包内容")
    add_table(doc, [
        ("星海知枢/", "主题清单、theme.css 与深浅背景图。"),
        ("xinghai-workbench/", "插件清单、main.js、styles.css 与八张运行资源图。"),
        ("README / INSTALL", "产品说明与跨平台安装步骤。"),
        ("两份 DOCX", "安装指导和工作台操作手册。"),
    ])
    note(doc, "数据说明", "发布包不含 data.json。插件只在本地读取知识库，并在用户明确操作时创建或修改 Markdown；不上传知识库内容。")
    doc.add_page_break()
    heading(doc, "3. macOS 安装")
    steps(doc, [
        "解压 `星海知枢-Obsidian主题-v3.2.0.zip`。",
        "在 Finder 中打开知识库；看不到 `.obsidian` 时按 `Command + Shift + .`。",
        "把整个 `星海知枢` 文件夹复制到 `知识库/.obsidian/themes/星海知枢/`。",
        "把整个 `xinghai-workbench` 文件夹复制到 `知识库/.obsidian/plugins/xinghai-workbench/`。",
        "启动 Obsidian，在外观设置中选择“星海知枢”，并在社区插件设置中启用“星海知枢工作台”。",
    ])
    heading(doc, "4. Windows 安装")
    steps(doc, [
        "解压 `星海知枢-Obsidian主题-v3.2.0.zip`。",
        "在资源管理器中开启“查看 → 显示 → 隐藏的项目”。",
        "把整个 `星海知枢` 文件夹复制到 `知识库\\.obsidian\\themes\\星海知枢\\`。",
        "把整个 `xinghai-workbench` 文件夹复制到 `知识库\\.obsidian\\plugins\\xinghai-workbench\\`。",
        "启动 Obsidian，在外观设置中选择“星海知枢”，并在社区插件设置中启用“星海知枢工作台”。",
    ])
    note(doc, "Windows 注意", "不需要管理员权限；不要把主题或插件复制到 Obsidian 程序目录或 `%APPDATA%`。", True)
    doc.add_page_break()
    heading(doc, "5. 首次启用与验收")
    bullets(doc, [
        "点击左侧 Ribbon 的主页图标，或运行命令“打开星海知枢工作台”。",
        "按提示设置任务、知识、文章和复盘等内容目录映射。",
        "验证新增任务、新增项目、本周复盘和 25/50 分钟专注功能。",
        "收起左右侧栏，确认左上角和右上角分别出现可执行的展开入口。",
        "切换深色与浅色模式，确认背景、文字、文件树、标签页和弹窗清晰可读。",
    ])
    heading(doc, "6. 升级、卸载与回滚")
    steps(doc, [
        "完全退出 Obsidian，并备份旧插件目录中的 data.json。",
        "用新版主题和插件程序文件覆盖对应目录，同时保留原 data.json。",
        "卸载时先停用插件并切换其他主题，再删除两个安装目录。",
        "重新启动 Obsidian，确认主题 3.2.0 与插件 1.2.6 状态正常。",
    ])
    note(doc, "笔记安全", "卸载不会自动删除已创建的 Markdown；如需清理笔记，必须由用户确认具体文件后操作。", True)
    doc.save(DOCS / "星海知枢-macOS与Windows安装指导说明.docx")


def build_manual():
    doc = Document()
    configure(doc, "主题与工作台操作手册")
    title(doc, "主题与工作台操作手册", "星图、任务、项目、专注、日历与外观")
    heading(doc, "1. 产品边界")
    body(doc, "星海知枢由主题与工作台插件组成。主题统一 Obsidian 外观；插件在当前知识库中提供工作台、内容汇总和用户触发的 Markdown 写入。")
    add_table(doc, [
        ("主题 3.2.0", "深浅背景、全局组件外观和主题清单。"),
        ("工作台 1.2.6", "星图、任务、项目、专注、日历、时间线、关联笔记和标签。"),
        ("数据权限", "本地读取知识库；仅在用户明确操作时写入 Markdown；不上传内容。"),
        ("不包含", "用户 data.json、测试库、QA 截图、归档数据和第三方依赖。"),
    ])
    heading(doc, "2. 工作台结构")
    add_table(doc, [
        ("目录星图", "以保留目录为节点，点击后打开对应入口笔记或文件夹。"),
        ("今日工作台", "显示翻页时钟、当前专注、进行中的项目、今日三件事和最近笔记。"),
        ("右侧信息", "显示本地日历、今日时间线、关联笔记和标签汇总。"),
        ("窄屏布局", "使用星图、工作台、日历、关联标签页重排内容。"),
    ])
    heading(doc, "3. 首次配置")
    steps(doc, [
        "启用插件后打开“星海知枢工作台”。",
        "在首次提示或插件设置中配置任务、知识、文章、灵感和复盘目录。",
        "检查目录映射是否指向当前知识库内已存在或允许创建的文件夹。",
        "完成后返回工作台，确认星图、统计和右侧信息能够读取本地内容。",
    ])
    heading(doc, "4. 任务、项目与复盘")
    bullets(doc, [
        "“新增任务”把任务写入指定目录；勾选“今日三件事”时同步写入当日日记。",
        "“新增项目”创建项目笔记，进行中的项目按笔记任务完成比例计算进度。",
        "“本周复盘”打开或创建当前周的复盘笔记。",
        "空内容、缺失目录或创建失败时，界面应给出提示而不是生成假数据。",
    ])
    heading(doc, "5. 专注计时")
    steps(doc, [
        "在当前专注卡片选择 25 分钟或 50 分钟。",
        "可直接使用默认目标，也可从项目或任务行把内容设为专注目标。",
        "点击“开始专注”；再次点击可结束并累计当天专注时长。",
        "重新加载工作区后，插件根据本地配置恢复仍有效的计时状态。",
    ])
    heading(doc, "6. 日历、时间线与关联信息")
    bullets(doc, [
        "日历只读取知识库中的日记与任务，不连接外部日历。",
        "点击没有日记的日期时，必须确认后才创建 Markdown 日记。",
        "今日时间线解析日记中带时间的事项。",
        "关联笔记与标签根据本地链接和元数据生成，点击项目可打开对应笔记。",
    ])
    heading(doc, "7. 左右侧栏")
    bullets(doc, [
        "侧栏展开时使用 Obsidian 原生收起按钮，不额外显示重复控制。",
        "左侧栏收起后，窗口左上方显示“展开左侧栏”入口。",
        "右侧栏收起后，窗口右上方显示“展开右侧栏”入口。",
        "左右同时收起时两个入口同时显示；点击后只展开对应侧栏并自动隐藏补位入口。",
    ])
    heading(doc, "8. 深色与浅色模式")
    steps(doc, [
        "打开 Obsidian“设置 → 外观”。",
        "在“基础颜色方案”中选择深色或浅色。",
        "确认正文、标题、链接、文件树和弹窗的对比度正常。",
    ])
    bullets(doc, [
        "深色模式使用紫蓝星海背景和深色半透明面板。",
        "浅色模式使用雾蓝星云背景和浅色高可读面板。",
        "两套模式包含在同一主题中，无需重复安装。",
    ])
    heading(doc, "9. 升级与卸载")
    steps(doc, [
        "备份 `.obsidian/appearance.json` 和插件目录中的 data.json。",
        "完全退出 Obsidian。",
        "升级时覆盖主题和插件程序文件，但保留原 data.json。",
        "卸载时先停用插件、切换其他主题，再删除两个安装目录。",
        "重新启动 Obsidian，检查主题、插件和已有 Markdown 笔记。",
    ])
    doc.add_page_break()
    heading(doc, "10. 常见问题")
    add_table(doc, [
        ("主题列表中没有出现", "检查目录层级，`manifest.json` 必须直接位于 `星海知枢` 文件夹。"),
        ("插件列表中没有出现", "确认 main.js、manifest.json、styles.css 位于 xinghai-workbench 根目录，并重新加载 Obsidian。"),
        ("工作台没有内容", "在插件设置中重新检查内容目录映射。"),
        ("侧栏展开入口不出现", "确认插件版本为 1.2.6，完全退出并重新打开 Obsidian。"),
        ("背景图片不显示", "确认 `assets` 中两个 starfield 文件存在且文件名未改变。"),
        ("更新后仍是旧样式", "完全退出 Obsidian 后覆盖文件，再重新启动。"),
        ("局部颜色异常", "暂时停用 CSS 片段或其他修改外观的组件后复核。"),
        ("需要恢复默认", "停用工作台插件，并在外观设置中切换为 Obsidian 默认主题。"),
    ])
    note(doc, "数据说明", "停用或删除插件不会自动删除已创建的笔记、附件、标签或 Properties；清理内容必须由用户确认具体文件。")
    doc.save(DOCS / "星海知枢-产品操作手册.docx")


if __name__ == "__main__":
    DOCS.mkdir(exist_ok=True)
    build_install()
    build_manual()
    print(DOCS / "星海知枢-macOS与Windows安装指导说明.docx")
    print(DOCS / "星海知枢-产品操作手册.docx")
