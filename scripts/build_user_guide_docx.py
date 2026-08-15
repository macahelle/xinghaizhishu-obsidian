from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "星海知枢使用与维护指导手册.docx"
LOGO = ROOT / "xinghai-workbench" / "assets" / "xinghai-logo-reference.png"

FONT = "Arial Unicode MS"
ACCENT = "6F4BC6"
ACCENT_DARK = "41297F"
INK = "202333"
MUTED = "62677A"
PALE = "F0ECFA"
PALE_BLUE = "E8EEF5"
LIGHT = "F7F8FB"
WHITE = "FFFFFF"
CAUTION = "FFF4D6"
CAUTION_INK = "775900"
RISK = "FCE8E8"
RISK_INK = "8A1C1C"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=None, bold=None, color=INK, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D8DBE5", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, ACCENT_DARK, 18, 10),
        "Heading 2": (13, ACCENT, 14, 7),
        "Heading 3": (12, ACCENT_DARK, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name, left, first, after in (
        ("List Bullet", 0.375, -0.188, 4),
        ("List Number", 0.375, -0.188, 4),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(left)
        style.paragraph_format.first_line_indent = Inches(first)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("星海知枢 · 使用与维护指导手册")
    set_run_font(hr, size=9, bold=True, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("候选版 3.2.0 / 插件 1.2.3   ·   第 ")
    set_run_font(fr, size=9, color=MUTED)
    add_page_number(fp)
    fr2 = fp.add_run(" 页")
    set_run_font(fr2, size=9, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        shape = p.add_run().add_picture(str(LOGO), width=Inches(0.78))
        shape._inline.docPr.set("descr", "星海知枢品牌 Logo")
        shape._inline.docPr.set("title", "星海知枢")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("星海知枢")
    set_run_font(r, size=30, bold=True, color=ACCENT_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("Obsidian 工作台使用与维护指导手册")
    set_run_font(r, size=17, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("适用于主题 3.2.0 / 插件 1.2.3")
    set_run_font(r, size=11, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(58)
    r = p.add_run("更新日期：2026-08-15")
    set_run_font(r, size=10, color=MUTED)

    add_callout(
        doc,
        "发布状态",
        "当前为候选版。桌面深浅主题、窄屏布局、侧栏收起/恢复和核心数据链路已完成测试；真实移动端设备截图与触控回归尚未完成。用户确认前不得覆盖正式 Tmac 知识库。",
        fill=CAUTION,
        title_color=CAUTION_INK,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Machaelle × Codex")
    set_run_font(r, size=10, color=MUTED)
    doc.add_page_break()


def add_callout(doc, title, body, fill=PALE, title_color=ACCENT_DARK):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    set_table_borders(table, color=fill, size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=title_color)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    widths = [2700, 6660]
    set_table_geometry(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, value in enumerate(("项目", "说明")):
        set_cell_shading(header.cells[i], PALE_BLUE)
        p = header.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(value)
        set_run_font(r, size=10.5, bold=True, color=ACCENT_DARK)
    for label, value in rows:
        cells = table.add_row().cells
        for index, text in enumerate((label, value)):
            p = cells[index].paragraphs[0]
            r = p.add_run(text)
            set_run_font(r, size=10.2, bold=index == 0)
            set_cell_margins(cells[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if len(table.rows) % 2 == 1:
            set_cell_shading(cells[0], LIGHT)
            set_cell_shading(cells[1], LIGHT)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_three_column_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for i, text in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], PALE_BLUE)
        r = table.rows[0].cells[i].paragraphs[0].add_run(text)
        set_run_font(r, size=10.2, bold=True, color=ACCENT_DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            r = cells[i].paragraphs[0].add_run(text)
            set_run_font(r, size=9.8, bold=i == 0)
            set_cell_margins(cells[i])
        if len(table.rows) % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, LIGHT)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_chapter_intro(doc, number, title, summary):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{number:02d}")
    set_run_font(r, size=11, bold=True, color=ACCENT)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(title)
    set_run_font(r, size=22, bold=True, color=ACCENT_DARK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(summary)
    set_run_font(r, size=11, color=MUTED)


def add_contents(doc):
    add_chapter_intro(doc, 0, "阅读导航", "按角色选择章节；普通用户优先阅读第 1–6 章，维护人员继续阅读第 7–10 章。")
    add_key_value_table(doc, [
        ("快速开始", "第 1 章：安装、启用、首次打开和安装后检查。"),
        ("界面认识", "第 2 章：三栏结构、动态星图、四模块和右侧信息栏。"),
        ("首次配置", "第 3 章：日记目录、周复盘目录、专注时长和内容映射。"),
        ("日常使用", "第 4–6 章：内容采集、今日三件事、项目、专注、日历和窄屏。"),
        ("维护升级", "第 7–9 章：数据规范、升级回滚、测试与发布门禁。"),
        ("故障排查", "第 10 章：常见现象、定位方法和恢复操作。"),
    ])
    add_callout(doc, "事实边界", "本手册以当前源码、测试库和候选包为准。V3.1 发布包已确认落后于源码，禁止继续安装或作为验证依据。")
    doc.add_page_break()


def build_manual(doc):
    add_cover(doc)
    add_contents(doc)

    add_chapter_intro(doc, 1, "快速安装与首次启动", "在不覆盖现有知识库数据的前提下安装主题和本地插件。")
    doc.add_heading("1.1 安装前准备", level=2)
    add_bullet(doc, "Obsidian 版本不低于 1.9.0；当前桌面实机验证版本为 1.13.7。")
    add_bullet(doc, "完全退出 Obsidian，备份知识库中的 .obsidian/workspace.json、graph.json，以及现有主题和插件目录。")
    add_bullet(doc, "首次只在独立测试库安装。正式 Tmac 知识库必须等移动端门禁和用户视觉确认完成后再处理。")
    doc.add_heading("1.2 安装文件", level=2)
    add_key_value_table(doc, [
        ("星海知枢/", "主题目录，放入 <知识库>/.obsidian/themes/星海知枢。"),
        ("xinghai-workbench/", "本地插件目录，放入 <知识库>/.obsidian/plugins/xinghai-workbench。"),
        ("工作台/00-星海工作台.md", "可选原生工作台内容，建议放入知识库入口目录。"),
        ("工作台/03-工作台视图.base", "可选最近笔记 Base 视图。"),
        ("workspace-xinghai.json", "可选桌面布局预设；替换前必须备份并退出 Obsidian。"),
        ("graph-xinghai.json", "可选图谱分组与颜色预设。"),
    ])
    doc.add_heading("1.3 启用步骤", level=2)
    for text in (
        "把主题与插件文件夹复制到对应 .obsidian 目录。",
        "打开 Obsidian，在“设置 → 外观 → 主题”选择“星海知枢”。",
        "进入“设置 → 第三方插件”，启用第三方插件并打开“星海知枢工作台”。",
        "点击左侧 Ribbon 的主页图标，或从命令面板运行“打开星海知枢工作台”。",
        "确认左侧文件树、中央星图与工作台、右侧信息栏同时可见且互不遮挡。",
    ):
        add_number(doc, text)
    add_callout(doc, "不要直接覆盖工作区", "插件本身不要求替换 workspace.json。只有确需套用预设时，才在完全退出 Obsidian并完成备份后替换。", fill=CAUTION, title_color=CAUTION_INK)
    add_chapter_intro(doc, 2, "认识工作台界面", "工作台使用 Obsidian 原生分栏，不通过跨栏浮层伪造页面。")
    add_key_value_table(doc, [
        ("左侧文件树", "保留知识库真实目录和 Obsidian 原生文件操作。"),
        ("中央上方星图", "读取全部一级目录动态生成节点；优先识别入口目录作为中心，节点可打开真实目录或笔记。"),
        ("中央下方工作台", "包含当前专注、进行中的项目、今日三件事、最近笔记四个模块。"),
        ("右侧信息栏", "显示日历、今日时间线、关联笔记和标签。"),
        ("顶部品牌入口", "Logo 与“星海知枢”组合位于中央栏顶部；点击返回已有工作台，不创建重复视图。"),
        ("左侧 Ribbon 主页入口", "在打开二级笔记后，一键返回原工作台。"),
        ("底部汇总", "显示全库反向链接、Properties、词数和字符数；只对变更文件重新统计。"),
    ])
    doc.add_heading("2.1 星图行为", level=2)
    add_bullet(doc, "一级目录变化会触发刷新；外围目录 1–9 个使用单环，10 个以上自动使用双环。")
    add_bullet(doc, "暗面星球属于装饰层，不对应目录，不会产生错误入口或可访问性噪声。")
    add_bullet(doc, "星链为动态曲线；目录数量、名称、坐标和连线索引没有写死。")
    doc.add_heading("2.2 深浅主题", level=2)
    add_bullet(doc, "进入“设置 → 外观 → 基础颜色方案”切换深色或浅色。")
    add_bullet(doc, "两个模式使用同一主题文件，Logo、中心星球、暗面星球和背景分别提供适配材质。")
    add_callout(doc, "视觉异常时", "先确认主题确实为“星海知枢”，再执行应用重载。不要通过修改知识库内容来修复纯视觉问题。")
    add_chapter_intro(doc, 3, "首次配置", "让工作台识别你的日记、复盘和内容保存目录。")
    doc.add_heading("3.1 设置页", level=2)
    add_key_value_table(doc, [
        ("日记目录", "今日三件事、日历和时间线的数据目录。示例：30-复盘沉淀/日记。"),
        ("周复盘目录", "“本周复盘”的保存与打开目录。示例：30-复盘沉淀/周复盘。"),
        ("默认专注时长", "可选择 25 或 50 分钟，工作台也可即时切换。"),
        ("内容写入映射", "扫描全库目录、文件名、Properties、标签和链接关系，为各类记录建议保存目录。"),
    ])
    doc.add_heading("3.2 内容映射原则", level=2)
    add_bullet(doc, "任务、灵感和文章是必需映射；用户可以暂时忽略首次提示，但创建未映射类型时必须手动选择真实目录。")
    add_bullet(doc, "插件不会把 00/10/20 等目录写死为保存目标；当前测试库路径只是扫描后的结果。")
    add_bullet(doc, "勾选“设为该类型默认目录”后，新映射写入插件 data.json。")
    add_bullet(doc, "密钥、账号和外部 API 不参与本插件配置；日历不连接外部服务。")
    add_callout(doc, "配置检查", "完成设置后创建一条测试记录，确认 Markdown 文件实际出现在选定目录，并检查 Properties 中的 captureType、targetFolder 和 created。")
    add_chapter_intro(doc, 4, "新增内容与知识采集", "所有记录写入真实 Markdown 和 Properties，不使用 HTML 假数据。")
    add_three_column_table(doc, ("内容类型", "写入类型", "典型用途"), [
        ("偶发灵感记录", "type: idea", "闪念、收件箱内容、待整理想法。"),
        ("工作任务", "type: task", "待办事项、项目行动；可加入今日三件事。"),
        ("外部文章知识采集", "type: article", "文章摘要、来源链接、知识提炼。"),
        ("主题知识笔记", "type: knowledge-note", "原子笔记、主题知识和长期积累。"),
        ("复盘沉淀记录", "type: review", "日记、回顾和经验沉淀。"),
        ("自定义记录", "type: capture", "不属于上述分类的临时内容。"),
    ], [2600, 2100, 4660])
    doc.add_heading("4.1 创建步骤", level=2)
    for text in (
        "点击“新增任务”或运行命令“新增今日任务”。",
        "在“记录内容”中输入多行文字：首个非空行作为标题，其余内容写入正文。",
        "选择内容类型和保存目录；外部文章可填写来源链接，任务可填写时间。",
        "工作任务可选择“加入今日三件事”；当日最多三项。",
        "提交后查看成功提示，并在文件树中确认新 Markdown 文件。",
    ):
        add_number(doc, text)
    doc.add_heading("4.2 错误恢复", level=2)
    add_bullet(doc, "内容为空时创建按钮保持禁用。")
    add_bullet(doc, "目录不存在、时间格式错误或提交失败时，错误在弹窗内显示，并把焦点移回对应字段。")
    add_bullet(doc, "若任务已保存但日记未配置或今日三件事已满，插件会明确提示“已保存但未加入今日”。")
    add_chapter_intro(doc, 5, "任务、项目与专注", "把知识采集转化为可追踪的日常行动。")
    doc.add_heading("5.1 今日三件事", level=2)
    add_bullet(doc, "数据位于当日日记的“## 今日三件事”标题下，读取前三条原生 Markdown 任务。")
    add_bullet(doc, "勾选时显示保存状态并禁用重复操作；成功后即时更新，失败时恢复原状态并提示。")
    add_bullet(doc, "如果日记不存在，创建任务流程会按配置生成 YYYY-MM-DD.md，并写入 type: daily 和 date。")
    doc.add_heading("5.2 进行中的项目", level=2)
    add_bullet(doc, "项目笔记需包含 Properties：type: project、status: active。")
    add_bullet(doc, "项目进度按正文中已完成任务数 ÷ 全部任务数自动计算，不维护静态百分比。")
    add_bullet(doc, "点击项目打开真实笔记；点击专注图标可把该项目设为当前专注目标。")
    doc.add_heading("5.3 专注计时", level=2)
    for text in (
        "选择 25 分钟或 50 分钟。",
        "确认当前专注目标后点击“开始专注”。",
        "工作台持续显示剩余时间；关闭或重载工作台后，会话状态保存在插件 data.json。",
        "结束后当日累计时长自动更新；不足一分钟按秒显示。",
    ):
        add_number(doc, text)
    add_callout(doc, "计时数据", "当前会话、默认时长、专注目标和每日累计秒数保存在 .obsidian/plugins/xinghai-workbench/data.json。停用主题不会删除这些数据。")
    add_chapter_intro(doc, 6, "日历、关联与窄屏操作", "在有限宽度下保持完整功能，并保留可逆的布局操作。")
    doc.add_heading("6.1 日历与时间线", level=2)
    add_bullet(doc, "点击已有日记的日期会直接打开文件。")
    add_bullet(doc, "点击没有日记的日期会先询问，只有确认后才创建文件，浏览不会静默写入。")
    add_bullet(doc, "今日时间线解析日记中带 HH:mm 的任务或列表项。")
    doc.add_heading("6.2 关联笔记与标签", level=2)
    add_bullet(doc, "关联图结合当前文件的入链、出链、入口特征和更新时间评分。")
    add_bullet(doc, "标签按钮显示真实使用次数和来源目录；点击后进入对应内容。")
    doc.add_heading("6.3 窄屏与侧栏恢复", level=2)
    add_bullet(doc, "中央容器小于 520px 时，自动切换为“星图 / 工作台 / 日历 / 关联”四标签单列布局。")
    add_bullet(doc, "标签支持鼠标、Tab 和方向键操作，并提供 selected 状态。")
    add_bullet(doc, "极窄时点击“收起侧栏”，插件记录左右栏各自的原始状态后再收起。")
    add_bullet(doc, "收起后，工作台左上角常驻“恢复侧栏”；点击后只恢复原本展开的栏，不会误打开原本关闭的栏。")
    add_callout(doc, "移动端状态", "移动端代码采用同一四标签结构，但真实设备上的触控、软键盘、安全区和截图尚未完成验收，当前不得视为正式发布通过。", fill=CAUTION, title_color=CAUTION_INK)
    add_chapter_intro(doc, 7, "数据结构与备份", "理解哪些信息在 Markdown 中，哪些信息属于插件状态。")
    add_key_value_table(doc, [
        ("普通知识内容", "保存在知识库 Markdown 文件和 Properties 中，可由 Obsidian 原生管理。"),
        ("今日三件事", "当日日记“## 今日三件事”下的原生任务。"),
        ("时间线", "当日日记中带 HH:mm 的任务或列表项。"),
        ("项目进度", "项目笔记正文任务的完成比例。"),
        ("内容映射", "插件 data.json 中的 contentMappings。"),
        ("专注状态", "插件 data.json 中的 focus、focusTarget、focusSecondsByDay。"),
        ("侧栏恢复状态", "插件 data.json 中的 sidebarRestoreState；恢复完成后回写为 null。"),
        ("工作区布局", ".obsidian/workspace.json；可选预设替换前必须单独备份。"),
    ])
    doc.add_heading("7.1 建议备份范围", level=2)
    add_bullet(doc, "整个知识库 Markdown 文件。")
    add_bullet(doc, ".obsidian/plugins/xinghai-workbench/data.json。")
    add_bullet(doc, ".obsidian/workspace.json、graph.json 和主题/插件目录。")
    add_bullet(doc, "升级前保留当前可运行版本压缩包，确保可以回滚。")
    add_chapter_intro(doc, 8, "升级、停用与回滚", "所有升级先在测试库完成，不跨版本覆盖未经核验的正式库。")
    doc.add_heading("8.1 升级流程", level=2)
    for text in (
        "完全退出 Obsidian，并备份主题、插件、data.json 和工作区文件。",
        "在独立测试库替换主题与 xinghai-workbench 插件目录。",
        "打开测试库，确认插件版本、工作台加载、内容写入、项目进度、专注持久化和侧栏恢复。",
        "完成深色、浅色、目标视口和窄屏截图；移动端发布还需真实设备证据。",
        "用户确认后再制定正式知识库安装窗口；当前候选版不执行此步。",
    ):
        add_number(doc, text)
    doc.add_heading("8.2 停用", level=2)
    add_bullet(doc, "停用主题：Markdown 和插件数据不会删除，仅恢复其他主题外观。")
    add_bullet(doc, "停用插件：动态星图、计时、写回和右栏停止显示，普通 Markdown 文件仍保留。")
    doc.add_heading("8.3 回滚", level=2)
    for text in (
        "完全退出 Obsidian。",
        "恢复备份的主题、插件目录及 data.json。",
        "如曾替换预设，恢复 workspace.json 和 graph.json。",
        "重新打开 Obsidian并选择原主题，检查文件树与笔记内容。",
    ):
        add_number(doc, text)
    add_callout(doc, "禁止事项", "不要安装 V3.1 包；该包已确认与当前源码和测试库漂移，缺少当前资产。只能使用经过四方哈希一致性验证的 V3.2.0 候选包。", fill=RISK, title_color=RISK_INK)
    add_chapter_intro(doc, 9, "开发维护与发布门禁", "供维护人员验证源码、测试库、候选目录和 ZIP 始终一致。")
    doc.add_heading("9.1 目录职责", level=2)
    add_key_value_table(doc, [
        ("星海知枢/", "主题源码。"),
        ("xinghai-workbench/", "插件源码、manifest、样式和视觉资产。"),
        ("test-vault/", "独立验收知识库；用户确认前只在此处安装。"),
        ("assets/reference/", "视觉唯一事实源。"),
        ("qa/", "局部、整屏、交互和发布证据。"),
        ("release/", "候选目录和压缩包；必须由构建脚本生成。"),
        ("scripts/build-release.js", "生成候选目录与干净 ZIP。"),
        ("tests/validate-package.js", "检查 JSON、CSS、路径、哈希、资产清单和 ZIP 内容。"),
    ])
    doc.add_heading("9.2 最小验证命令", level=2)
    add_bullet(doc, "node --check xinghai-workbench/main.js")
    add_bullet(doc, "node tests/plugin-core.test.js")
    add_bullet(doc, "node scripts/build-release.js")
    add_bullet(doc, "node tests/validate-package.js")
    add_bullet(doc, "unzip -t 'release/星海知枢-工作台-v3.2.0-候选版.zip'")
    doc.add_heading("9.3 发布通过条件", level=2)
    add_bullet(doc, "自动测试、JSON/CSS/路径检查全部通过。")
    add_bullet(doc, "源码、测试库、候选目录和 ZIP 的文件哈希及资产清单一致。")
    add_bullet(doc, "深色、浅色、1229 × 768 和窄屏同屏截图通过，无 P0/P1/P2。")
    add_bullet(doc, "真实移动端截图与触控回归通过。")
    add_bullet(doc, "项目记忆已更新，用户确认视觉方向。")
    add_callout(doc, "当前状态", "截至 2026-08-15，桌面门禁与插件 1.2.3 品牌对齐复测通过；真实移动端设备证据仍缺失，因此只能保留候选状态。", fill=CAUTION, title_color=CAUTION_INK)
    add_chapter_intro(doc, 10, "常见问题排查", "先读取现象和日志，再按最小范围修复，不反复猜测。")
    add_three_column_table(doc, ("现象", "优先检查", "处理方式"), [
        ("工作台没有出现", "插件是否启用；版本是否为 1.2.3", "运行“打开星海知枢工作台”；必要时重载 Obsidian。"),
        ("只有背景没有功能", "是否只安装了主题", "同时安装并启用 xinghai-workbench 插件。"),
        ("新增内容不能保存", "内容、目标目录、时间格式和映射", "按弹窗内错误修正；确认目录存在且可写。"),
        ("今日三件事不显示", "日记目录和标题结构", "检查日记目录，并确保存在“## 今日三件事”和原生任务。"),
        ("项目进度不更新", "Properties 和正文任务", "确认 type: project、status: active，并使用 Markdown 任务。"),
        ("点击日期创建了文件", "日期原先是否无日记且已确认", "无日记日期只有确认后才创建；若未确认仍创建，记录复现步骤。"),
        ("窄屏内容被挤压", "中央容器是否小于 520px", "切换顶部标签；需要时点击“收起侧栏”。"),
        ("侧栏收起后无法返回", "左上角“恢复侧栏”是否可见", "使用恢复入口；若缺失，确认插件版本并检查 sidebarRestoreState。"),
        ("Logo 或标题错位", "最终渲染局部截图", "不要只看元素盒；测量可见像素边界和加权视觉重心。"),
        ("安装包与测试效果不同", "四方哈希与资产清单", "重新运行构建和 validate-package，禁止手工拼包。"),
    ], [2350, 3000, 4010])
    doc.add_heading("10.1 提交问题时应提供", level=2)
    add_bullet(doc, "Obsidian 版本、主题版本、插件版本和操作系统。")
    add_bullet(doc, "问题发生前后的完整操作步骤。")
    add_bullet(doc, "同一视口截图；视觉问题同时提供局部放大图。")
    add_bullet(doc, "控制台或日志中的关键错误，不粘贴无关长日志。")
    add_bullet(doc, "是否发生在 test-vault，正式知识库是否保持未修改。")
    doc.add_heading("10.2 验收清单", level=2)
    for text in (
        "主题和插件正常加载，版本正确。",
        "主页入口能从二级笔记返回原工作台。",
        "任务写回、项目进度和专注计时持久化正常。",
        "日历无静默创建，时间线读取真实日记。",
        "窄屏标签与侧栏收起/恢复形成完整闭环。",
        "深浅主题和品牌视觉锚点均有当前证据。",
        "候选包哈希、资产清单和 ZIP 完整性通过。",
        "移动端未验证时，发布状态保持候选。",
    ):
        add_bullet(doc, text)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— 手册结束 —")
    set_run_font(r, size=10, color=MUTED)


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    build_manual(doc)
    props = doc.core_properties
    props.title = "星海知枢 Obsidian 工作台使用与维护指导手册"
    props.subject = "安装、配置、使用、维护、故障排查与发布验收"
    props.author = "Machaelle × Codex"
    props.keywords = "Obsidian, 星海知枢, 工作台, 使用手册, 维护指南"
    props.comments = "基于主题 3.2.0 / 插件 1.2.3 候选版编制"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
