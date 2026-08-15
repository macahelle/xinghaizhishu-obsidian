from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCS.mkdir(exist_ok=True)

BLUE = "635BFF"
BLUE_DARK = "302B73"
CYAN = "31C6D4"
TEXT = "202536"
MUTED = "667085"
LIGHT = "EEF1FF"
PALE = "F6F7FC"
WHITE = "FFFFFF"
RED = "B42318"
FONT = "Arial Unicode MS"
EAST_ASIA_FONT = "Arial Unicode MS"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def fix_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, size=None, bold=None, color=None, italic=None, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def paragraph_border_bottom(paragraph, color=BLUE, size="10", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for node in (fld_begin, instr, fld_sep, text, fld_end):
        run._r.append(node)
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc, short_title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, BLUE_DARK, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"星海知枢  |  {short_title}")
    set_run_font(r, size=8.5, bold=True, color=MUTED)
    paragraph_border_bottom(p, color="D9DCEF", size="5", space="4")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("主题 3.2.0  ·  插件 1.2.4  ·  第 ")
    set_run_font(r, size=8.5, color=MUTED)
    add_page_field(p)
    r = p.add_run(" 页")
    set_run_font(r, size=8.5, color=MUTED)


def add_cover(doc, title, subtitle, audience):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(68)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("星海知枢")
    set_run_font(r, size=14, bold=True, color=CYAN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    set_run_font(r, size=28, bold=True, color=BLUE_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run(subtitle)
    set_run_font(r, size=13, color=MUTED)

    table = doc.add_table(rows=4, cols=2)
    fix_table_geometry(table, [2700, 6660])
    set_repeat_table_header(table.rows[0])
    rows = [
        ("适用版本", "主题 3.2.0 / 插件 1.2.4"),
        ("运行环境", "macOS 或 Windows + Obsidian 1.9.0 或更高版本"),
        ("适用对象", audience),
        ("更新日期", "2026-08-15"),
    ]
    for idx, (label, value) in enumerate(rows):
        left, right = table.rows[idx].cells
        set_cell_shading(left, LIGHT)
        for cell, text, bold, color in ((left, label, True, BLUE_DARK), (right, value, False, TEXT)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=9.5, bold=bold, color=color)

    add_callout(doc, "发布说明", "移动端已完成代码适配、窄容器模拟和自动检测；未在真实手机 Obsidian 中完成截图与触控回归。", warning=True)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=TEXT)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, color=TEXT)
    else:
        r = p.add_run(text)
        set_run_font(r, color=TEXT)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_run_font(r, color=TEXT)


def create_numbering_definition(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_steps(doc, items):
    num_id = create_numbering_definition(doc)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.25
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num])
        p_pr.insert(0, num_pr)
        r = p.add_run(item)
        set_run_font(r, color=TEXT)


def add_callout(doc, label, text, warning=False, trailing_space=True):
    table = doc.add_table(rows=1, cols=1)
    fix_table_geometry(table, [9360])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF2F0" if warning else LIGHT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"{label}：")
    set_run_font(r, size=9.5, bold=True, color=RED if warning else BLUE_DARK)
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=TEXT)
    if trailing_space:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(1)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    fix_table_geometry(table, [2700, 6660])
    headers = table.rows[0].cells
    for cell, text in zip(headers, ("项目", "说明")):
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_run_font(r, size=9.5, bold=True, color=BLUE_DARK)
    set_repeat_table_header(table.rows[0])
    for label, value in rows:
        cells = table.add_row().cells
        for cell, text, bold in ((cells[0], label, True), (cells[1], value, False)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=9.2, bold=bold, color=TEXT)
    fix_table_geometry(table, [2700, 6660])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure(doc, path, caption, width=6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    picture = run.add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption.split("  ", 1)[-1])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(caption)
    set_run_font(r, size=8.5, italic=True, color=MUTED)


def add_section_page(doc, number, title, intro):
    if len(doc.paragraphs) > 1:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{number:02d}")
    set_run_font(r, size=11, bold=True, color=CYAN)
    add_heading(doc, title, 1)
    add_body(doc, intro)


def build_installation_guide():
    doc = Document()
    configure_document(doc, "macOS / Windows 安装指导")
    add_cover(doc, "macOS / Windows 安装指导说明", "从备份、安装到首次验收的完整流程", "macOS 与 Windows 首次安装者、知识库管理员")

    add_heading(doc, "快速导航", 1)
    add_key_value_table(doc, [
        ("1. 安装前", "核对 Obsidian 版本，备份 `.obsidian` 和重要 Markdown。"),
        ("2. 复制文件", "把主题和插件复制到指定的隐藏目录。"),
        ("3. 启用", "在 Obsidian 中启用主题和“星海知枢工作台”插件。"),
        ("4. 验收", "检查三栏、星图、四模块、日历与写回反馈。"),
        ("5. 回滚", "如果异常，停用插件并恢复备份。"),
    ])
    add_callout(doc, "推荐", "首次安装先使用知识库副本，验证无误后再安装到正式库。")

    add_section_page(doc, 1, "安装前准备", "星海知枢是 Obsidian 内的主题与本地插件，不是独立桌面应用。macOS 不需要 `.dmg`/`.pkg`，Windows 不需要 `.exe`/`.msi`；两个系统使用同一份 ZIP。")
    add_key_value_table(doc, [
        ("最低版本", "Obsidian 1.9.0"),
        ("已验证桌面版", "Obsidian 1.13.7（macOS）"),
        ("主题版本", "3.2.0"),
        ("插件版本", "1.2.4"),
        ("正式包", "星海知枢-工作台-v3.2.0-正式版.zip"),
    ])
    add_heading(doc, "1.1 必须备份", 2)
    add_bullets(doc, [
        "整个知识库中的 Markdown 文件。",
        "`.obsidian/workspace.json` 和 `.obsidian/graph.json`。",
        "`.obsidian/plugins/xinghai-workbench/data.json`（如果已安装旧版）。",
        "现有 `.obsidian/themes` 与 `.obsidian/plugins/xinghai-workbench` 目录。",
    ])
    add_callout(doc, "注意", "复制工作区预设前必须完全退出 Obsidian，否则应用退出时可能用旧状态覆盖新文件。", warning=True)

    add_section_page(doc, 2, "识别安装包", "解压正式 ZIP 后，核心内容包括主题、插件、可选工作区预设和说明文件。")
    add_key_value_table(doc, [
        ("`星海知枢/`", "主题文件夹，含 `theme.css`、`manifest.json` 和视觉资产。"),
        ("`xinghai-workbench/`", "本地插件，含 `main.js`、`styles.css`、`manifest.json` 和资产。"),
        ("`工作台/`", "桌面工作区、外观和图谱参考预设。"),
        ("`README.md`", "产品概览、结构与兼容说明。"),
        ("`INSTALL.md`", "简版安装、启用和回滚说明。"),
    ])

    add_section_page(doc, 3, "macOS 与 Windows 完整安装", "以知识库 `Tmac` 为例；请把示例路径替换为你自己的实际知识库路径。两个系统的库内相对目录完全相同。")
    add_heading(doc, "3.1 知识库路径示例", 2)
    add_key_value_table(doc, [
        ("macOS", "`/Users/用户名/Documents/Tmac`，主题放到 `Tmac/.obsidian/themes/星海知枢/`。"),
        ("Windows", "`D:\\Obsidian\\Tmac` 或 `C:\\Users\\用户名\\Documents\\Tmac`，主题放到 `Tmac\\.obsidian\\themes\\星海知枢\\`。"),
        ("共同插件路径", "知识库根目录下的 `.obsidian/plugins/xinghai-workbench/`。"),
    ])
    add_heading(doc, "3.2 复制与启用", 2)
    add_steps(doc, [
        "完全退出 Obsidian，完成上一章的备份。",
        "解压 `星海知枢-工作台-v3.2.0-正式版.zip`。",
        "把整个 `星海知枢` 文件夹复制到 `Tmac/.obsidian/themes/星海知枢/`。",
        "把整个 `xinghai-workbench` 文件夹复制到 `Tmac/.obsidian/plugins/xinghai-workbench/`。",
        "启动 Obsidian，进入“设置 → 外观 → 主题”，选择“星海知枢”。",
        "进入“设置 → 第三方插件”，允许第三方插件并启用“星海知枢工作台”。",
        "点击左侧 Ribbon 中的主页图标，或在命令面板运行“打开星海知枢工作台”。",
    ])
    add_callout(doc, "隐藏目录", "macOS Finder 中看不到 `.obsidian` 时，在知识库文件夹中按 `Command + Shift + .`。Windows 资源管理器中若不显示，进入“查看 → 显示 → 隐藏的项目”，或直接在地址栏输入 `.obsidian`。")
    add_callout(doc, "Windows 注意", "本包不含独立可执行文件，正常安装不需要管理员权限。请不要把插件复制到 Obsidian 程序安装目录或 `%APPDATA%`；必须放在当前知识库的 `.obsidian` 内。")

    add_section_page(doc, 4, "首次启动与验收", "启用后不要只检查背景，还要确认插件视图和真实数据链路。")
    add_bullets(doc, [
        "左侧是 Obsidian 原生文件树，中央显示知识星图与今日工作台，右侧显示日历、时间线、关联笔记和标签。",
        "星图应按知识库的实际一级目录动态生成，节点可点击。",
        "今日三件事的勾选应写回当日 Markdown 日记，并显示保存状态。",
        "项目进度应根据项目笔记的任务完成比自动计算。",
        "“进行中的项目”右上角应显示新增按钮，打开表单后由用户选择保存目录。",
        "专注计时支持 25 和 50 分钟，刷新后会话与当日累计数据应保留。",
        "点击无日记的日期应先询问，不得因浏览静默新建文件。",
    ])
    add_callout(doc, "验收限制", "移动端仅完成代码适配、窄容器模拟和自动检测；未完成真实手机触控验收。", warning=True)

    add_section_page(doc, 5, "可选预设、升级与回滚", "主题和插件可独立运行，不强制覆盖用户的原有工作区。")
    add_heading(doc, "5.1 可选预设", 2)
    add_bullets(doc, [
        "`workspace-xinghai.json`：桌面工作区参考布局。",
        "`graph-xinghai.json`：图谱分组与配色参考。",
        "`appearance-xinghai.json`：主题启用配置参考。",
    ])
    add_body(doc, "使用任何预设前，先备份同名原文件，并在 Obsidian 完全退出后替换。")
    add_heading(doc, "5.2 升级", 2)
    add_steps(doc, [
        "备份当前主题、插件和 `data.json`。",
        "先在知识库副本中替换主题与插件目录。",
        "验证插件加载、任务写回、项目进度、专注持久化和工作区恢复。",
        "确认无误后再更新正式知识库。",
    ])
    add_heading(doc, "5.3 回滚", 2)
    add_steps(doc, [
        "完全退出 Obsidian。",
        "恢复备份的主题、插件、`data.json`、`workspace.json` 和 `graph.json`。",
        "重新启动，选择原主题并检查笔记内容。",
    ])

    add_section_page(doc, 6, "常见问题", "先核对路径、版本和插件状态，再根据现象处理。")
    add_key_value_table(doc, [
        ("主题不出现", "检查 `Tmac/.obsidian/themes/星海知枢/theme.css` 是否存在，避免多套一层文件夹。"),
        ("插件不出现", "检查 `plugins/xinghai-workbench/manifest.json` 与 `main.js`，然后重启 Obsidian。"),
        ("显示已安装但无界面", "确认已启用插件，使用 Ribbon 主页图标或命令面板打开工作台。"),
        ("今日三件事无数据", "检查日记目录映射，以及当日日记中是否存在 `## 今日三件事`。"),
        ("布局拥挤", "在窄容器中使用顶部四标签，或点击“收起侧栏”。"),
        ("需要完全卸载", "停用插件和主题，再删除对应的主题/插件文件夹；Markdown 笔记不会被自动删除。"),
        ("Windows 解压后路径过长", "先把 ZIP 解压到较短目录，例如 `D:\\Xinghai`，再复制主题和插件文件夹。"),
    ])

    path = DOCS / "星海知枢-macOS与Windows安装指导说明.docx"
    doc.save(path)
    return path


def build_product_manual():
    doc = Document()
    configure_document(doc, "产品操作手册")
    add_cover(doc, "产品操作手册", "从首次配置到日常知识与行动管理", "日常使用者、知识库管理员、产品维护者")

    add_heading(doc, "阅读导航", 1)
    add_key_value_table(doc, [
        ("新用户", "先阅读第 1–4 章：产品定位、界面、首次配置和内容采集。"),
        ("日常使用者", "重点阅读第 5–7 章：任务、项目、专注、日历、关联与窄屏。"),
        ("维护者", "阅读第 8–10 章：数据结构、备份、排障和验收。"),
    ])

    add_section_page(doc, 1, "产品定位与数据原则", "星海知枢是运行在 Obsidian 内的个人知识与行动工作台，不是只有背景的静态主题。")
    add_key_value_table(doc, [
        ("主题层", "提供深色星海、浅色星雾以及 Obsidian 全局组件外观。"),
        ("插件层", "提供星图、时钟、统计、项目、任务、专注、日历、时间线、关联和标签。"),
        ("工作区层", "保留 Obsidian 原生左侧文件树、中央工作台与右侧信息面板。"),
        ("数据层", "任务、项目、复盘和内容继续保存为 Markdown/Properties，不用 HTML 假数据。"),
    ])
    add_callout(doc, "核心原则", "工作台可停用，但你的知识仍是普通 Markdown；主题和插件不应锁定数据。")

    add_section_page(doc, 2, "认识工作台", "桌面布局采用 Obsidian 原生分栏，主要信息在同一屏内展示。")
    screenshot = ROOT / "qa/user-ux-remediation-2026-08-15/03-workbench-dark.png"
    if screenshot.exists():
        add_figure(doc, screenshot, "图 1  深色主题下的星海知枢三栏工作台")
    add_heading(doc, "2.1 左侧文件树", 2)
    add_body(doc, "保留你的原生目录结构；星图会读取全部一级目录，目录数量变化后自动重新布局。")
    add_heading(doc, "2.2 中央星图与今日工作台", 2)
    add_bullets(doc, [
        "知识星图：优先识别入口目录为中心，其余一级目录按单环或双环布局。",
        "今日工作台：翻页时钟、当前专注、进行中的项目、今日三件事和最近笔记。",
        "主操作：“新增任务”和“本周复盘”。",
    ])
    add_heading(doc, "2.3 右侧信息面板", 2)
    add_body(doc, "包含日历、今日时间线、关联笔记图和标签。所有统计从知识库实时读取。")

    add_section_page(doc, 3, "首次配置与目录映射", "插件不把 `00/10/20` 等目录名写死，而是扫描当前知识库后给出建议。")
    add_steps(doc, [
        "打开工作台或插件设置，查看日记、周复盘与内容类型的目录建议。",
        "核对任务、灵感和文章的保存目录；这三类是必需映射。",
        "根据你的知识库设置主题知识、复盘和自定义记录目录。",
        "创建内容时可勾选“设为该类型默认目录”，新映射保存在插件 `data.json`。",
    ])
    add_key_value_table(doc, [
        ("日记", "今日三件事、时间线和日历数据来源。"),
        ("周复盘", "“本周复盘”按钮的打开/创建目标。"),
        ("任务", "工作任务记录，可同步加入今日三件事。"),
        ("灵感", "偶发灵感、闪念与临时捕捉。"),
        ("文章", "外部文章、摘要、来源链接和自动关键词。"),
    ])
    add_callout(doc, "保护规则", "缺少映射时不会静默回退到硬编码目录；创建时必须明确选择实际保存位置。")

    add_section_page(doc, 4, "新增内容与知识采集", "内容会写入真实 Markdown 和 Properties，首个非空行作为标题，其余内容保留在正文。")
    add_steps(doc, [
        "点击“新增任务”，或从命令面板运行“新增今日任务”。",
        "在宽版多行输入区输入标题和说明。",
        "选择偶发灵感、工作任务、外部文章、主题知识、复盘或自定义记录。",
        "选择保存目录；任务可设置时间并选择是否加入今日三件事。",
        "提交后查看成功或失败提示，再在文件树中确认新笔记。",
    ])
    add_heading(doc, "4.1 表单反馈", 2)
    add_bullets(doc, [
        "内容为空时，“创建”按钮保持禁用。",
        "目录不存在或时间格式错误时，弹窗内显示错误并将焦点移回问题字段。",
        "任务已保存但日记未配置，或今日三件事已满时，插件会说明未加入原因。",
    ])

    add_section_page(doc, 5, "今日三件事、项目与专注", "工作台把 Markdown 中的行动数据聚合为当日可执行的工作界面。")
    add_heading(doc, "5.1 今日三件事", 2)
    add_bullets(doc, [
        "数据位于当日日记 `## 今日三件事` 下的前三条原生任务。",
        "勾选时会显示操作中状态，并禁止重复点击。",
        "保存成功后即时更新；失败时恢复原状态并提示。",
    ])
    add_heading(doc, "5.2 进行中的项目", 2)
    add_body(doc, "项目笔记需要 Properties：`type: project` 与 `status: active`。进度 = 正文中已完成任务数 ÷ 全部任务数；没有任务时显示 0%。")
    add_steps(doc, [
        "点击“进行中的项目”卡片右上角的新增按钮。",
        "首行填写项目名称，其余内容填写项目目标。",
        "选择保存目录；若该类型尚未映射，可选择是否记住该目录。",
        "点击“创建记录”，插件写入 `type: project`、`status: active` 的 Markdown 项目笔记并刷新项目列表。",
    ])
    add_heading(doc, "5.3 专注计时", 2)
    add_steps(doc, [
        "选择 25 分钟或 50 分钟。",
        "核对当前专注目标，点击“开始专注”。",
        "在工作台查看剩余时间和当日累计。",
        "重载或重启后，插件从 `data.json` 恢复会话状态。",
    ])

    add_section_page(doc, 6, "日历、时间线、关联与标签", "右侧面板只读取 Obsidian 内部数据，不连接外部日历或网络服务。")
    add_key_value_table(doc, [
        ("日历", "按 macOS 月历的周日到周六顺序显示完整六周、今天、选中日期以及有无日记状态。"),
        ("时间线", "解析当日日记中带 `HH:mm` 的任务或列表项。"),
        ("关联笔记", "结合当前笔记的入链、出链、入口特征和更新时间生成关联图。"),
        ("标签", "显示真实标签使用次数和来源目录，点击后进入相关内容。"),
    ])
    add_callout(doc, "防误写", "点击已有日记的日期会直接打开；点击无日记的日期必须先确认，不会因浏览静默创建 Markdown。")

    add_section_page(doc, 7, "窄屏与移动布局", "当中央工作台容器小于 520px 时，布局自动切换为标签式单列。")
    screenshot = ROOT / "qa/user-ux-remediation-2026-08-15/04-narrow-workbench-fixed.png"
    if screenshot.exists():
        add_figure(doc, screenshot, "图 2  桌面窄容器模拟：四标签与单列工作台", width=5.5)
    add_bullets(doc, [
        "顶部提供“星图 / 工作台 / 日历 / 关联”四个标签。",
        "标签支持点击、Tab、左右方向键、Home 和 End。",
        "极窄状态可点击“收起侧栏”；操作后使用左上角常驻“恢复侧栏”还原。",
        "星图外围节点会限制在容器边界内，长目录名自动省略。",
    ])
    add_callout(
        doc,
        "未验证项",
        "3.2.0 只完成移动代码适配、窄容器模拟和自动断言；未在真实手机 Obsidian 执行触控回归。",
        warning=True,
        trailing_space=False,
    )

    add_section_page(doc, 8, "数据、备份与隐私", "工作台把业务内容保存为 Markdown，仅把设置和运行状态保存在插件 `data.json`。")
    add_key_value_table(doc, [
        ("Markdown/Properties", "任务、项目、日记、复盘、灵感、文章和主题知识。"),
        ("`data.json`", "内容映射、日记/复盘路径、当前专注会话、当日专注累计和侧栏恢复状态。"),
        ("`workspace.json`", "Obsidian 的桌面分栏与打开视图状态。"),
        ("主题/插件目录", "可执行文件、样式和图像资产，升级前建议保留当前可运行副本。"),
    ])
    add_bullets(doc, [
        "插件不要求 API Key，不连接外部日历。",
        "密钥、账号、数据库密码不应写入模板、插件源码或 Markdown 示例。",
        "备份时优先备份整个知识库，至少包含上表四类数据。",
    ])

    add_section_page(doc, 9, "停用、升级与回滚", "主题、插件和 Markdown 数据相互分离，尽量使故障可逆。")
    add_heading(doc, "9.1 停用", 2)
    add_bullets(doc, [
        "停用主题只会更换外观，不删除 Markdown 或插件数据。",
        "停用插件会停止动态星图、计时、写回和右栏视图，普通笔记仍然存在。",
    ])
    add_heading(doc, "9.2 升级", 2)
    add_steps(doc, [
        "备份当前主题、插件、`data.json` 和工作区文件。",
        "先在知识库副本中替换文件，不跨版本直接覆盖唯一正式库。",
        "验证加载、写回、项目进度、专注持久化、工作区恢复与深浅主题。",
    ])
    add_heading(doc, "9.3 回滚", 2)
    add_steps(doc, [
        "完全退出 Obsidian。",
        "恢复备份的主题、插件目录、`data.json`、`workspace.json` 和 `graph.json`。",
        "重新打开知识库，检查文件树、笔记、任务和工作台入口。",
    ])

    add_section_page(doc, 10, "常见问题与验收清单", "出现问题时先读取现象和日志，再按最小范围修复，不反复猜测。")
    add_heading(doc, "10.1 常见问题", 2)
    add_key_value_table(doc, [
        ("星图缺少目录", "确认目录位于知识库一级，重新打开工作台；不要在代码里手工添加节点。"),
        ("项目不显示", "检查 Properties 是否为 `type: project` 和 `status: active`。"),
        ("进度不对", "检查项目正文中的 `- [ ]` 和 `- [x]` 原生任务，不要手写静态百分比。"),
        ("今日三件事写回失败", "查看弹出提示，检查日记目录和标题。失败时界面应恢复原勾选状态。"),
        ("窄屏裁切", "确认中央容器已切换四标签；如仍太窄，先使用“收起侧栏”。"),
    ])
    add_heading(doc, "10.2 交付验收清单", 2)
    add_bullets(doc, [
        "主题 3.2.0 和插件 1.2.4 正常加载。",
        "Ribbon 主页入口能从二级笔记返回原工作台。",
        "任务写回、项目进度、专注持久化与工作区恢复正常。",
        "日历不静默创建，时间线读取真实日记。",
        "深色、浅色、1229 × 768 与窄容器模拟证据都使用当前版本。",
        "正式 ZIP 的包结构、文件哈希和资产清单验签通过。",
        "文档明确标注移动端未完成真机触控回归。",
    ])

    path = DOCS / "星海知枢-产品操作手册.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    outputs = [build_installation_guide(), build_product_manual()]
    for output in outputs:
        print(output.relative_to(ROOT))
